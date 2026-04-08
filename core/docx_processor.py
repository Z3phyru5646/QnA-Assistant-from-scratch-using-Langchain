"""
DOCX Processor — Extract text, images, and tables from .docx files.
Returns same format as PDFProcessor for pipeline compatibility.
"""

import logging
from pathlib import Path
from config.settings import DATA_DIR

logger = logging.getLogger(__name__)


class DocxProcessor:
    """Splits a DOCX into text blocks, images, and tables."""

    def __init__(self):
        self.images_dir = DATA_DIR / "extracted_images"
        self.tables_dir = DATA_DIR / "extracted_tables"
        self.images_dir.mkdir(parents=True, exist_ok=True)
        self.tables_dir.mkdir(parents=True, exist_ok=True)

    def process(self, docx_path: Path):
        """
        Process a DOCX file.
        Returns: (text_blocks, image_paths, table_data)
        """
        docx_path = Path(docx_path)
        source_name = docx_path.name
        logger.info(f"Processing DOCX: {source_name}")

        text_blocks = []
        image_paths = []
        table_data = []

        try:
            from docx import Document as DocxDocument
            import pandas as pd

            doc = DocxDocument(str(docx_path))

            # Extract text paragraphs
            current_text = []
            page_num = 1  # DOCX doesn't have real pages, approximate

            for para in doc.paragraphs:
                if para.text.strip():
                    current_text.append(para.text.strip())

                # Approximate page break every ~3000 chars
                combined = '\n'.join(current_text)
                if len(combined) > 3000:
                    text_blocks.append({
                        "text": combined,
                        "page": page_num,
                        "source": source_name,
                    })
                    current_text = []
                    page_num += 1

            # Remaining text
            if current_text:
                text_blocks.append({
                    "text": '\n'.join(current_text),
                    "page": page_num,
                    "source": source_name,
                })

            # Extract images
            img_idx = 0
            for rel in doc.part.rels.values():
                if "image" in str(rel.reltype):
                    try:
                        image_data = rel.target_part.blob
                        if len(image_data) > 5120:  # Skip tiny icons
                            ext = rel.target_part.content_type.split('/')[-1]
                            if ext == 'jpeg':
                                ext = 'jpg'
                            img_filename = f"{docx_path.stem}_img{img_idx + 1}.{ext}"
                            img_path = self.images_dir / img_filename
                            with open(img_path, "wb") as f:
                                f.write(image_data)
                            image_paths.append({
                                "path": str(img_path),
                                "page": 1,
                                "source": source_name,
                            })
                            img_idx += 1
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")

            # Extract tables
            for tbl_idx, table in enumerate(doc.tables):
                try:
                    rows_data = []
                    for row in table.rows:
                        rows_data.append([cell.text.strip() for cell in row.cells])

                    if len(rows_data) > 1:
                        headers = rows_data[0]
                        df = pd.DataFrame(rows_data[1:], columns=headers)

                        csv_filename = f"{docx_path.stem}_tbl{tbl_idx + 1}.csv"
                        csv_path = self.tables_dir / csv_filename
                        df.to_csv(csv_path, index=False)

                        try:
                            md = df.to_markdown(index=False)
                        except Exception:
                            md = df.to_string(index=False)

                        table_data.append({
                            "dataframe": df,
                            "page": 1,
                            "source": source_name,
                            "markdown": md,
                            "csv_path": str(csv_path),
                        })
                except Exception as e:
                    logger.warning(f"Failed to process table {tbl_idx}: {e}")

        except ImportError:
            logger.error("python-docx is not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")

        logger.info(
            f"DOCX extraction — Text: {len(text_blocks)} blocks, "
            f"Images: {len(image_paths)}, Tables: {len(table_data)}"
        )
        return text_blocks, image_paths, table_data
